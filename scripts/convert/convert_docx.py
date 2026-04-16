#!/usr/bin/env python3
"""Convert Word .docx files to Markdown.

Usage:
    python convert_docx.py <input_file> <output_md> [--images-dir <dir>]

Exit code 0 on success, non-zero on failure.
Prints JSON summary to stdout on completion.
"""

import argparse
import io
import json
import os
import re
import sys
import uuid
from pathlib import Path
from typing import Dict, List

import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image


# ---------------------------------------------------------------------------
# Header / footer detection
# ---------------------------------------------------------------------------

def is_header_or_footer(paragraph, all_paragraphs_text: List[str]) -> bool:
    """Detect whether a paragraph is likely a header or footer."""
    text = paragraph.text.strip()
    if not text:
        return False

    # 1. Pure page numbers
    if text.isdigit() and len(text) <= 4:
        return True

    # 2. Common page-number formats
    page_patterns = [
        r'^第\s*\d+\s*页$',
        r'^-\s*\d+\s*-$',
        r'^\d+\s*/\s*\d+$',
        r'^Page\s+\d+$',
        r'^\d+\s+of\s+\d+$',
    ]
    for pattern in page_patterns:
        if re.match(pattern, text, re.IGNORECASE):
            return True

    # 3. Repeated short text (likely running header)
    if len(text) < 100:
        count = all_paragraphs_text.count(text)
        if count > 3:
            return True

    # 4. Common footer keywords
    footer_keywords = [
        '版权所有', '保留所有权利', 'all rights reserved', 'copyright',
        '机密', 'confidential', '内部资料',
    ]
    text_lower = text.lower()
    if any(kw in text_lower for kw in footer_keywords) and len(text) < 150:
        return True

    return False


# ---------------------------------------------------------------------------
# Heading helpers
# ---------------------------------------------------------------------------

def get_heading_level(style_name: str) -> int:
    """Return the Markdown heading level for a Word style name."""
    mapping = {
        'Heading 1': 1, 'Title': 1,
        'Heading 2': 2, 'Subtitle': 2,
        'Heading 3': 3,
        'Heading 4': 4,
        'Heading 5': 5,
        'Heading 6': 6,
    }
    return mapping.get(style_name, 1)


# ---------------------------------------------------------------------------
# Paragraph processing
# ---------------------------------------------------------------------------

def process_paragraph(paragraph) -> List[str]:
    """Convert a single DOCX paragraph to a list of Markdown lines."""
    lines: List[str] = []
    text = paragraph.text.strip()

    if not text:
        return [""]

    style_name = paragraph.style.name if paragraph.style else ""

    # Headings via Word style
    if style_name.startswith('Heading'):
        level = get_heading_level(style_name)
        lines.append(f"{'#' * level} {text}")
        lines.append("")
    elif style_name.lower() in ('title', 'subtitle'):
        lines.append(f"# {text}")
        lines.append("")
    else:
        alignment = paragraph.alignment
        if alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            # Centre-aligned short text → treat as heading
            if len(text) < 50 and '。' not in text and '.' not in text:
                lines.append(f"## {text}")
            else:
                lines.append(f"**{text}**")
        else:
            # List detection
            if re.match(r'^\s*[\d\w]+[.)]\s+', text) or text.strip().startswith(('•', '-', '*')):
                if re.match(r'^\s*\d+[.)]\s+', text):
                    match = re.match(r'^\s*\d+[.)]\s+(.+)', text)
                    if match:
                        lines.append(f"1. {match.group(1)}")
                    else:
                        lines.append(text)
                else:
                    cleaned = re.sub(r'^\s*[•\-*]\s+', '', text)
                    lines.append(f"- {cleaned}")
            else:
                lines.append(text)
        lines.append("")

    return lines


# ---------------------------------------------------------------------------
# Table conversion
# ---------------------------------------------------------------------------

def convert_table_to_markdown(table) -> List[str]:
    """Convert a python-docx Table object to Markdown table lines."""
    md: List[str] = []
    try:
        rows = table.rows
        if not rows:
            return []

        has_content = any(
            cell.text.strip() for row in rows for cell in row.cells
        )
        if not has_content:
            return ["*[Empty table]*"]

        header_row = rows[0]
        headers = [
            cell.text.strip() or f"Col{i+1}"
            for i, cell in enumerate(header_row.cells)
        ]
        md.append("| " + " | ".join(headers) + " |")
        md.append("| " + " | ".join(["---"] * len(headers)) + " |")

        for row in rows[1:]:
            cells = [
                cell.text.strip().replace('\n', '<br>') or "-"
                for cell in row.cells
            ]
            md.append("| " + " | ".join(cells) + " |")
    except Exception:
        return ["*[Table parsing failed]*"]

    return md


# ---------------------------------------------------------------------------
# Image helpers
# ---------------------------------------------------------------------------

def get_image_extension(image_data: bytes) -> str:
    """Determine image format from raw bytes."""
    try:
        img = Image.open(io.BytesIO(image_data))
        fmt = (img.format or "png").lower()
        return "jpg" if fmt == "jpeg" else fmt.strip()
    except Exception:
        return "png"


def validate_image(path: Path) -> bool:
    """Return True if *path* is a valid image file."""
    try:
        with Image.open(path) as img:
            img.verify()
        return True
    except Exception:
        return False


def extract_images(doc, images_dir: Path) -> List[Dict]:
    """Extract all images from the document relationships and save them."""
    images_dir.mkdir(parents=True, exist_ok=True)
    images_info: List[Dict] = []

    try:
        for rel_id, rel in doc.part.rels.items():
            if "image" not in rel.target_ref:
                continue
            try:
                image_data = rel.target_part.blob
                if len(image_data) < 100:
                    continue

                ext = get_image_extension(image_data)
                filename = f"{uuid.uuid4()}.{ext}"
                dest = images_dir / filename

                with open(dest, 'wb') as f:
                    f.write(image_data)

                if validate_image(dest):
                    images_info.append({
                        "filename": filename,
                        "path": str(dest),
                        "size": len(image_data),
                        "format": ext,
                        "rel_id": rel_id,
                    })
                else:
                    dest.unlink(missing_ok=True)
            except Exception:
                continue
    except Exception:
        pass

    return images_info


def get_paragraph_images(paragraph, all_images: List[Dict]) -> List[Dict]:
    """Return images embedded inside a specific paragraph."""
    result: List[Dict] = []
    try:
        for run in paragraph.runs:
            if 'graphic' not in run._element.xml:
                continue
            ns_blip = (
                './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'
                '//{http://schemas.openxmlformats.org/drawingml/2006/picture}blipFill'
                '//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
            )
            for blip in run._element.findall(ns_blip):
                embed = blip.get(
                    '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                )
                if embed:
                    for img in all_images:
                        if img.get('rel_id') == embed:
                            result.append(img)
                            break
    except Exception:
        pass
    return result


# ---------------------------------------------------------------------------
# Markdown cleanup
# ---------------------------------------------------------------------------

def clean_markdown(lines: List[str]) -> str:
    """Join *lines* into a single Markdown string, trimming excess blanks."""
    content = "\n".join(lines)
    content = re.sub(r'\n{3,}', '\n\n', content)
    stripped = [l.rstrip() for l in content.split('\n')]
    while stripped and not stripped[0].strip():
        stripped.pop(0)
    while stripped and not stripped[-1].strip():
        stripped.pop()
    return '\n'.join(stripped)


# ---------------------------------------------------------------------------
# Main conversion
# ---------------------------------------------------------------------------

def convert_docx(input_file: str, output_md: str, images_dir: str) -> Dict:
    """Convert *input_file* (.docx) → Markdown written to *output_md*.

    Returns a summary dict.
    """
    input_path = Path(input_file)
    output_path = Path(output_md)
    img_dir = Path(images_dir)

    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    doc = docx.Document(str(input_path))

    total_paragraphs = len(doc.paragraphs)
    total_tables = len(doc.tables)

    # Collect paragraph texts for header/footer detection
    all_paragraphs_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Extract images
    images_info = extract_images(doc, img_dir)

    # Compute the relative path from the output markdown file to the images dir
    try:
        img_rel = os.path.relpath(img_dir, output_path.parent)
    except ValueError:
        # On Windows, relpath fails across drives
        img_rel = str(img_dir)

    markdown_content: List[str] = []
    table_index = 0
    processed_tables: set = set()

    # Walk elements in document-body order to preserve interleaving
    for element in doc.element.body:
        # --- tables ---
        if element.tag.endswith('tbl'):
            if table_index < len(doc.tables):
                table = doc.tables[table_index]
                if table_index not in processed_tables:
                    try:
                        table_md = convert_table_to_markdown(table)
                        if table_md:
                            markdown_content.append("")
                            markdown_content.extend(table_md)
                            markdown_content.append("")
                        processed_tables.add(table_index)
                    except Exception:
                        markdown_content.append("\n*[Table parsing failed]*\n")
                table_index += 1

        # --- paragraphs ---
        elif element.tag.endswith('p'):
            for paragraph in doc.paragraphs:
                if paragraph._element is not element:
                    continue

                # Filter headers / footers
                if is_header_or_footer(paragraph, all_paragraphs_text):
                    break

                md_lines = process_paragraph(paragraph)
                markdown_content.extend(md_lines)

                # Inline images
                p_images = get_paragraph_images(paragraph, images_info)
                for img in p_images:
                    rel_img_path = os.path.join(img_rel, img['filename']).replace('\\', '/')
                    markdown_content.append(f"![Image]({rel_img_path})")
                    markdown_content.append("")
                break

    result_md = clean_markdown(markdown_content)

    if not result_md.strip():
        result_md = "*Document content is empty or cannot be parsed*"

    # Write output
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(result_md, encoding='utf-8')

    return {
        "success": True,
        "file_type": "docx",
        "paragraphs": total_paragraphs,
        "tables": total_tables,
        "images": len(images_info),
    }


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> int:
    parser = argparse.ArgumentParser(
        description="Convert a .docx file to Markdown."
    )
    parser.add_argument("input_file", help="Path to the .docx file")
    parser.add_argument("output_md", help="Path to write the resulting .md file")
    parser.add_argument(
        "--images-dir",
        default=None,
        help="Directory to save extracted images (default: same dir as output_md)",
    )
    args = parser.parse_args()

    images_dir = args.images_dir
    if images_dir is None:
        images_dir = str(Path(args.output_md).parent)

    try:
        summary = convert_docx(args.input_file, args.output_md, images_dir)
        print(json.dumps(summary))
        return 0
    except Exception as exc:
        err = {"success": False, "error": str(exc)}
        print(json.dumps(err), file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())
